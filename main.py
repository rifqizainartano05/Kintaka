from flask import Flask, request, jsonify, render_template, send_file, redirect
from flask_cors import CORS
from groq import Groq
from docx import Document
from docx.shared import Inches

import numpy as np
import io
import os
import shutil
import tempfile
import subprocess
import threading
import time
import json
from datetime import datetime
from PIL import Image

from dotenv import load_dotenv
load_dotenv()

import firebase_admin
from firebase_admin import credentials, firestore, storage

app = Flask(__name__, template_folder="templates")
CORS(app, resources={r"/*": {"origins": "*"}}, expose_headers=["Content-Disposition"])

client = Groq(api_key=os.environ.get("GROQ_API_KEY", "YOUR_GROQ_API_KEY"))



# Initialize Firebase
try:
    bucket_name = os.environ.get('FIREBASE_STORAGE_BUCKET', 'kintaka-c1e74.firebasestorage.app')
    if not firebase_admin._apps:
        # Check if service account key exists
        if os.path.exists('firebase_credentials.json'):
            cred = credentials.Certificate('firebase_credentials.json')
            firebase_admin.initialize_app(cred, {
                'storageBucket': bucket_name
            })
            print("Firebase initialized with credentials file.")
        else:
            # Fallback to default credentials
            firebase_admin.initialize_app(options={
                'storageBucket': bucket_name
            })
            print("Firebase initialized with default credentials.")
    db = firestore.client()
    bucket = storage.bucket()
except Exception as e:
    print(f"Warning: Firebase failed to initialize. Make sure you set up firebase_credentials.json and FIREBASE_STORAGE_BUCKET. Error: {e}")
    db = None
    bucket = None

def log_history(filename: str, tool_used: str, status: str, local_path: str = ""):
    if db is None:
        print("Firebase is not initialized. Skipping history log.")
        return
    try:
        download_url = ""
        if status == "Completed" and local_path and os.path.exists(local_path) and bucket is not None:
            try:
                storage_path = f"history/{int(datetime.now().timestamp())}_{filename}"
                blob = bucket.blob(storage_path)
                blob.upload_from_filename(local_path)
                blob.make_public()
                download_url = blob.public_url
                print(f"File uploaded to Firebase Storage: {download_url}")
            except Exception as e:
                print(f"Failed to upload to Firebase Storage: {e}")

        date_processed = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc_ref = db.collection("history").document()
        doc_ref.set({
            "filename": filename,
            "tool_used": tool_used,
            "date_processed": date_processed,
            "status": status,
            "download_url": download_url,
            "timestamp": datetime.now()
        })
    except Exception as e:
        print(f"Error logging to Firebase: {e}")

def cleanup_temp_dir(dir_path: str):
    """Remove temporary directory"""
    try:
        if os.path.exists(dir_path):
            shutil.rmtree(dir_path)
    except Exception as e:
        print(f"Error cleaning up temp dir {dir_path}: {e}")

def delayed_cleanup(temp_dir: str):
    """Clean up temp directory after a short delay so send_file can finish reading it."""
    def cleanup():
        time.sleep(5)
        cleanup_temp_dir(temp_dir)
    threading.Thread(target=cleanup).start()

@app.route("/", methods=["GET"])
def root():
    return render_template("dashboard.html")

@app.route("/rupa-kata-page", methods=["GET"])
def rupa_kata_page():
    return render_template("rupa_kata.html")

@app.route("/alih-rupa-page", methods=["GET"])
def alih_rupa_page():
    return render_template("alih_rupa.html")

@app.route("/nalar-naskah-page", methods=["GET"])
def nalar_naskah_page():
    return render_template("nalar_naskah.html")

@app.route("/api/history", methods=["GET"])
def get_history():
    if db is None:
        return jsonify({"history": []})
    
    try:
        history_ref = db.collection("history").order_by("timestamp", direction=firestore.Query.DESCENDING).limit(100)
        docs = history_ref.stream()
        
        history_list = []
        for doc in docs:
            data = doc.to_dict()
            history_list.append({
                "id": doc.id,
                "filename": data.get("filename", ""),
                "tool_used": data.get("tool_used", ""),
                "date_processed": data.get("date_processed", ""),
                "status": data.get("status", ""),
                "download_url": data.get("download_url", "")
            })
        return jsonify({"history": history_list})
    except Exception as e:
        print(f"Error fetching history: {e}")
        return jsonify({"history": []}), 500

@app.route("/download-history/<item_id>", methods=["GET"])
def download_history(item_id):
    if db is None:
        return jsonify({"detail": "Firebase not initialized"}), 500
        
    try:
        doc_ref = db.collection("history").document(item_id)
        doc = doc_ref.get()
        
        if not doc.exists:
            return jsonify({"detail": "File not found"}), 404
            
        data = doc.to_dict()
        download_url = data.get("download_url")
        
        if download_url:
            return redirect(download_url)
        else:
            return jsonify({"detail": "No download URL available for this file"}), 404
            
    except Exception as e:
        return jsonify({"detail": f"Error downloading file: {e}"}), 500

@app.route("/ocr/", methods=["POST"])
def perform_ocr():
    if 'file' not in request.files:
        return jsonify({"detail": "No file part"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"detail": "No selected file"}), 400
    if not file.content_type.startswith('image/'):
        return jsonify({"detail": "Invalid file type. Please upload an image."}), 400
        
    output_format = request.form.get("output_format", "docx")
    
    try:
        contents = file.read()
        temp_dir = tempfile.mkdtemp()
        
        doc = Document()
        temp_image_path = os.path.join(temp_dir, "uploaded_image.png")
        
        img = Image.open(io.BytesIO(contents))
        img.save(temp_image_path, format="PNG")
            
        try:
            doc.add_picture(temp_image_path, width=Inches(6.0))
        except Exception as pic_err:
            print(f"Warning: Could not add image to docx: {pic_err}")
            
        out_filename = file.filename.rsplit('.', 1)[0] + "_converted.docx"
        docx_path = os.path.join(temp_dir, out_filename)
        doc.save(docx_path)
        
        if output_format.lower() == "pdf":
            import pythoncom
            from docx2pdf import convert
            pythoncom.CoInitialize()
            pdf_filename = file.filename.rsplit('.', 1)[0] + "_converted.pdf"
            pdf_path = os.path.join(temp_dir, pdf_filename)
            convert(docx_path, pdf_path)
            pythoncom.CoUninitialize()
            
            log_history(pdf_filename, "Rupa Kata (Image to Doc)", "Completed", pdf_path)
            delayed_cleanup(temp_dir)
            return send_file(pdf_path, as_attachment=True, download_name=pdf_filename, mimetype='application/pdf')
        else:
            log_history(out_filename, "Rupa Kata (Image to Doc)", "Completed", docx_path)
            delayed_cleanup(temp_dir)
            return send_file(docx_path, as_attachment=True, download_name=out_filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        
    except Exception as e:
        log_history(file.filename, "Rupa Kata (Image to Doc)", "Failed", "")
        return jsonify({"detail": f"Error processing image: {str(e)}"}), 500

@app.route("/docx2pdf/", methods=["POST"])
def convert_docx_to_pdf_endpoint():
    if 'file' not in request.files:
        return jsonify({"detail": "No file part"}), 400
    file = request.files['file']
    if file.filename == '' or not file.filename.endswith('.docx'):
        return jsonify({"detail": "Invalid file type. Please upload a .docx file."}), 400
        
    try:
        import pythoncom
        from docx2pdf import convert
        
        pythoncom.CoInitialize()
        
        temp_dir = tempfile.mkdtemp()
        docx_path = os.path.join(temp_dir, file.filename)
        pdf_filename = file.filename.rsplit('.', 1)[0] + ".pdf"
        pdf_path = os.path.join(temp_dir, pdf_filename)
        
        file.save(docx_path)
            
        convert(docx_path, pdf_path)
        pythoncom.CoUninitialize()
        
        log_history(pdf_filename, "Alih Rupa (Word to PDF)", "Completed", pdf_path)
        delayed_cleanup(temp_dir)
        return send_file(pdf_path, as_attachment=True, download_name=pdf_filename, mimetype='application/pdf')
        
    except Exception as e:
        log_history(file.filename, "Alih Rupa (Word to PDF)", "Failed", "")
        return jsonify({"detail": f"Error converting document: {str(e)}"}), 500

@app.route("/pdf2docx/", methods=["POST"])
def convert_pdf_to_docx():
    if 'file' not in request.files:
        return jsonify({"detail": "No file part"}), 400
    file = request.files['file']
    if file.filename == '' or not file.filename.lower().endswith('.pdf'):
        return jsonify({"detail": "Invalid file type. Please upload a PDF file."}), 400
    
    temp_dir = tempfile.mkdtemp()
    
    try:
        pdf_path = os.path.join(temp_dir, file.filename)
        file.save(pdf_path)
            
        from pdf2docx import Converter
        out_filename = file.filename.rsplit('.', 1)[0] + "_converted.docx"
        docx_path = os.path.join(temp_dir, out_filename)
        
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        log_history(out_filename, "Alih Rupa (PDF to Word)", "Completed", docx_path)
        delayed_cleanup(temp_dir)
        return send_file(docx_path, as_attachment=True, download_name=out_filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        
    except Exception as e:
        cleanup_temp_dir(temp_dir)
        log_history(file.filename, "Alih Rupa (PDF to Word)", "Failed", "")
        return jsonify({"detail": str(e)}), 500

@app.route("/ai-edit/", methods=["POST"])
def process_ai_edit():
    if 'file' not in request.files:
        return jsonify({"detail": "No file part"}), 400
    file = request.files['file']
    if file.filename == '' or not file.filename.endswith('.docx'):
        return jsonify({"detail": "Invalid file type. Please upload a .docx file."}), 400
        
    prompt = request.form.get("prompt")
    if not prompt:
        return jsonify({"detail": "No prompt provided"}), 400
        
    output_format = request.form.get("output_format", "docx")
    
    try:
        temp_dir = tempfile.mkdtemp()
        docx_path = os.path.join(temp_dir, file.filename)
        file.save(docx_path)
            
        doc = Document(docx_path)
        
        all_paragraphs = []
        def add_p(p):
            if p.text.strip() and p not in all_paragraphs:
                all_paragraphs.append(p)

        for p in doc.paragraphs:
            add_p(p)
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        add_p(p)
        
        input_data = []
        for i, p in enumerate(all_paragraphs):
            input_data.append({"id": i, "text": p.text.replace('\t', '[TAB]')})
        
        system_prompt = """You are an AI document editor. Process the document text according to the user's instruction.
CRITICAL RULES:
1. You will receive a JSON array of objects, each containing an 'id' and 'text'.
2. Apply the user's instruction (e.g. translation, correction, filling data) to the 'text' fields.
3. You MUST return a JSON array containing ONLY the objects that were modified.
4. Each returned object MUST have the original 'id' and the new modified 'text'.
5. If an object's text does not need to be changed, DO NOT include it in your output.
6. Preserve the exact number and placement of [TAB] markers in your modified text.
7. Return ONLY the raw JSON array. Do not include markdown wrappers."""
        
        user_message = f"Instruction: {prompt}\n\nDocument Data:\n{json.dumps(input_data, ensure_ascii=False)}"
        
        print(f"--- AI INSTRUCTION ---\n{prompt}")
        
        chat_completion = client.chat.completions.create(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message}
            ],
            model="llama-3.1-8b-instant",
            temperature=0.3,
        )
        
        result_text = chat_completion.choices[0].message.content.strip()
        
        if result_text.startswith('```'):
            result_text = "\n".join(result_text.split('\n')[1:])
            if result_text.endswith('```'):
                result_text = result_text[:-3]
        
        try:
            output_data = json.loads(result_text)
            if not isinstance(output_data, list):
                if isinstance(output_data, dict):
                    output_data = [output_data]
                    
            output_dict = {str(item["id"]): item["text"] for item in output_data if "id" in item and "text" in item}
            
            if not output_dict:
                return jsonify({"detail": "AI merasa tidak ada yang perlu diubah berdasarkan instruksi Anda. Coba perjelas instruksi Anda (contoh: 'Ubah nama menjadi Ahmad')."}), 400
            
            for i, p in enumerate(all_paragraphs):
                key = str(i)
                if key in output_dict:
                    new_text_val = output_dict[key]
                    
                    if isinstance(new_text_val, str):
                        new_text = new_text_val.replace('[TAB]', '\t')
                        old_text_raw = p.text
                        
                        if new_text != old_text_raw:
                            if ':' in old_text_raw and ':' in new_text:
                                old_prefix = old_text_raw.split(':', 1)[0]
                                if '\t' in old_prefix:
                                    clean_new_prefix = new_text.split(':', 1)[0].replace('\t', '').strip()
                                    new_suffix = new_text.split(':', 1)[1].strip()
                                    tab_count = old_prefix.count('\t')
                                    new_text = f"{clean_new_prefix}{chr(9) * tab_count}: {new_suffix}"
                                    
                            parts = new_text.split('\t')
                            
                            if p.runs:
                                for j in range(1, len(p.runs)):
                                    p.runs[j].text = ""
                                    
                                run = p.runs[0]
                                run.text = "" 
                                
                                for k, part in enumerate(parts):
                                    if part:
                                        run.add_text(part)
                                    if k < len(parts) - 1:
                                        run.add_tab()
                            else:
                                p.text = ""
                                run = p.add_run()
                                for k, part in enumerate(parts):
                                    if part:
                                        run.add_text(part)
                                    if k < len(parts) - 1:
                                        run.add_tab()
                                        
        except json.JSONDecodeError as e:
            print(f"JSON Parsing Error: {e}")
            return jsonify({"detail": "Sistem AI gagal memproses format teks. Silakan coba instruksi yang lebih sederhana."}), 500
            
        out_filename = file.filename.rsplit('.', 1)[0] + "_ai_edited.docx"
        out_path = os.path.join(temp_dir, out_filename)
        doc.save(out_path)
        
        media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        
        if output_format == "pdf":
            try:
                import pythoncom
                from docx2pdf import convert
                pythoncom.CoInitialize()
                
                pdf_filename = file.filename.rsplit('.', 1)[0] + "_ai_edited.pdf"
                pdf_path = os.path.join(temp_dir, pdf_filename)
                
                convert(out_path, pdf_path)
                
                out_path = pdf_path
                out_filename = pdf_filename
                media_type = 'application/pdf'
                pythoncom.CoUninitialize()
            except Exception as e:
                print(f"PDF Conversion failed: {e}")
                log_history(file.filename, "Nalar Naskah (AI)", "Failed", "")
                return jsonify({"detail": f"AI processing succeeded, but PDF conversion failed: {str(e)}"}), 500
        
        log_history(out_filename, "Nalar Naskah (AI)", "Completed", out_path)
        delayed_cleanup(temp_dir)
        return send_file(out_path, as_attachment=True, download_name=out_filename, mimetype=media_type)
        
    except Exception as e:
        log_history(file.filename, "Nalar Naskah (AI)", "Failed", "")
        return jsonify({"detail": f"Error during AI processing: {str(e)}"}), 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8000, debug=True)
